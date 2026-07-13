import base64
import csv
import re
from copy import copy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME_TYPE = "application/pdf"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
XLS_MIME_TYPE = "application/vnd.ms-excel"


@dataclass(frozen=True)
class ExcelEditResult:
    path: Path
    applied: list[str]
    skipped: list[str]
    verified: bool = False
    verification_errors: list[str] | None = None


def extract_document_text(path: Path, mime_type: str | None = None) -> str:
    suffix = path.suffix.lower()

    if suffix == ".pdf" or mime_type == "application/pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".xls" or mime_type == XLS_MIME_TYPE:
        return _extract_legacy_xls(path)
    if suffix == ".csv" or mime_type == "text/csv":
        return _extract_csv(path)
    if suffix in {".txt", ".md"} or (mime_type and mime_type.startswith("text/")):
        return path.read_text(encoding="utf-8", errors="replace")

    return ""


def build_image_content_block(path: Path, mime_type: str) -> dict[str, Any]:
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": mime_type,
            "data": encoded,
        },
    }


def clean_excel_workbook(source: Path, output_dir: Path, instruction_text: str = "") -> Path:
    """Create a lightly cleaned Excel copy without the source file."""
    source = prepare_excel_source(source, output_dir)
    workbook = load_workbook(str(source))
    heading = _extract_heading(instruction_text, source)
    for sheet in workbook.worksheets:
        _clean_sheet(sheet, heading, light=True)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    output_name = f"{source.stem}-cleaned-{timestamp}.xlsx"
    target = output_dir / output_name
    workbook.save(str(target))
    return target


def edit_excel_workbook(
    source: Path,
    output_dir: Path,
    instruction_text: str = "",
    plan: dict[str, Any] | None = None,
) -> ExcelEditResult:
    """Apply planned Excel edits, then return a formatted copy."""
    source = prepare_excel_source(source, output_dir)
    workbook = load_workbook(str(source))
    applied: list[str] = []
    skipped: list[str] = []

    actions = _actions_from_plan_and_instruction(plan, instruction_text)

    for sheet in workbook.worksheets:
        header_row = _find_header_row(sheet)
        for action in actions:
            action_type = str(action.get("type", "")).lower().strip()
            if action_type == "delete_columns":
                requested_columns = _string_list(action.get("columns"))
                deleted = _delete_requested_columns(sheet, header_row, requested_columns)
                if deleted:
                    applied.append(f"{sheet.title}: deleted column(s) {', '.join(deleted)}")
                elif requested_columns:
                    skipped.append(f"{sheet.title}: sikuipata column {', '.join(requested_columns)}")
            elif action_type == "keep_columns":
                requested_columns = _string_list(action.get("columns"))
                kept = _keep_requested_columns(sheet, header_row, requested_columns)
                if kept:
                    applied.append(f"{sheet.title}: kept only {', '.join(kept)}")
                elif requested_columns:
                    skipped.append(f"{sheet.title}: sikuipata columns za kubakiza")
            elif action_type == "rename_columns":
                renamed = _rename_requested_columns(sheet, header_row, action.get("columns"))
                if renamed:
                    applied.append(f"{sheet.title}: renamed {', '.join(renamed)}")
                else:
                    skipped.append(f"{sheet.title}: sikuweza kubadilisha header")
            elif action_type == "sort_by":
                sorted_by = _sort_sheet_by_column(
                    sheet,
                    header_row,
                    str(action.get("column") or ""),
                    str(action.get("direction") or "asc"),
                )
                if sorted_by:
                    applied.append(f"{sheet.title}: sorted by {sorted_by}")
                else:
                    skipped.append(f"{sheet.title}: sikuweza kusort column hiyo")

    heading = _extract_heading(instruction_text, source)
    if plan and str(plan.get("title") or "").strip():
        heading = _clean_heading(str(plan["title"]))

    has_data_actions = any(
        str(action.get("type", "")).lower().strip()
        in {"delete_columns", "keep_columns", "rename_columns", "sort_by", "add_product_summary"}
        for action in actions
    )

    for sheet in workbook.worksheets:
        _clean_sheet(sheet, heading, light=not has_data_actions)

    for sheet in workbook.worksheets:
        header_row = _find_header_row(sheet)
        for action in actions:
            if str(action.get("type", "")).lower().strip() == "add_product_summary":
                if _add_product_summary(sheet, header_row):
                    applied.append(f"{sheet.title}: added product summary")
                else:
                    skipped.append(f"{sheet.title}: sikuweza kutengeneza product summary")

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    output_name = f"{source.stem}-edited-{timestamp}.xlsx"
    target = output_dir / output_name
    workbook.save(str(target))
    verification_errors = verify_excel_result(target, {"actions": actions})
    return ExcelEditResult(
        path=target,
        applied=applied,
        skipped=skipped,
        verified=not verification_errors,
        verification_errors=verification_errors,
    )


def verify_excel_result(path: Path, plan: dict[str, Any] | None = None) -> list[str]:
    actions = _normalize_excel_plan_actions(plan)
    if not actions:
        return []
    errors: list[str] = []
    workbook = load_workbook(str(path), data_only=True)
    try:
        for action in actions:
            action_type = str(action.get("type", "")).lower().strip()
            if action_type == "delete_columns":
                columns = _string_list(action.get("columns"))
                for sheet in workbook.worksheets:
                    header_row = _find_header_row(sheet)
                    headers = _normalized_headers(sheet, header_row)
                    for column in columns:
                        requested = _normalize_header(column)
                        if requested and any(
                            requested == header or requested in header or header in requested
                            for header in headers
                        ):
                            errors.append(f"{sheet.title}: column bado ipo: {column}")
            elif action_type == "add_product_summary":
                if not any(_sheet_contains_value(sheet, "PRODUCT SUMMARY") for sheet in workbook.worksheets):
                    errors.append("Product summary haikuonekana kwenye output")
    finally:
        workbook.close()
    return errors


def _actions_from_plan_and_instruction(
    plan: dict[str, Any] | None,
    instruction_text: str,
) -> list[dict[str, Any]]:
    actions = _normalize_excel_plan_actions(plan)
    text_actions = _actions_from_instruction(instruction_text)

    merged: list[dict[str, Any]] = []
    for action in [*actions, *text_actions]:
        action_type = str(action.get("type", "")).lower().strip()
        if action_type == "delete_columns":
            columns = _string_list(action.get("columns"))
            existing = next(
                (
                    item
                    for item in merged
                    if str(item.get("type", "")).lower().strip() == "delete_columns"
                ),
                None,
            )
            if existing is not None:
                existing["columns"] = _unique_preserve_order(
                    _string_list(existing.get("columns")) + columns
                )
            elif columns:
                merged.append({"type": "delete_columns", "columns": columns})
            continue

        if action_type and not any(
            str(item.get("type", "")).lower().strip() == action_type for item in merged
        ):
            merged.append(action)
    return merged


def _actions_from_instruction(text: str) -> list[dict[str, Any]]:
    lowered = text.lower()
    actions: list[dict[str, Any]] = []
    delete_columns = _extract_delete_column_requests(text)

    if any(word in lowered for word in ["simu", "phone", "mobile", "namba", "contact"]):
        delete_columns.append("simu")
    if any(word in lowered for word in ["quantity", "quantiti", "qty", "qnty", "idadi", "pcs"]):
        delete_columns.append("quantity")
    if delete_columns:
        actions.append({"type": "delete_columns", "columns": _unique_preserve_order(delete_columns)})

    wants_summary = any(
        phrase in lowered
        for phrase in [
            "product summary",
            "summary chini",
            "summary iko chini",
            "item summary",
            "summary ya product",
            "muhtasari wa bidhaa",
        ]
    )
    if wants_summary:
        actions.append({"type": "add_product_summary"})

    return actions


def prepare_excel_source(source: Path, output_dir: Path) -> Path:
    """Convert legacy Excel formats into xlsx so the editor can safely write them."""
    suffix = source.suffix.lower()
    if suffix != ".xls":
        return source

    pd = _pandas()
    sheets = pd.read_excel(source, sheet_name=None, engine="xlrd")
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = output_dir / f"{source.stem}-converted-{timestamp}.xlsx"
    with pd.ExcelWriter(target, engine="openpyxl") as writer:
        for sheet_name, frame in sheets.items():
            safe_sheet_name = str(sheet_name)[:31] or "Sheet"
            frame.to_excel(writer, sheet_name=safe_sheet_name, index=False)
    return target


def excel_workbook_preview(source: Path, max_rows: int = 8) -> str:
    workbook = load_workbook(str(source), data_only=True, read_only=True)
    try:
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append(values)
                if len(rows) >= max_rows:
                    break
            if not rows:
                chunks.append(f"Sheet: {sheet.title}\n(empty)")
                continue
            header_index = _guess_header_index(rows)
            headers = rows[header_index] if header_index is not None else rows[0]
            sample_rows = rows[header_index + 1 : header_index + 4] if header_index is not None else rows[1:4]
            chunks.append(
                "\n".join(
                    [
                        f"Sheet: {sheet.title}",
                        "Likely headers: " + " | ".join(headers),
                        "Sample rows:",
                        *[" | ".join(row) for row in sample_rows],
                    ]
                )
            )
        return "\n\n".join(chunks)
    finally:
        workbook.close()


def is_specific_excel_edit_requested(text: str) -> bool:
    lowered = text.lower()
    return any(
        word in lowered
        for word in [
            "delete",
            "remove",
            "drop",
            "futa",
            "ondoa",
            "toa column",
            "toa safu",
            "rename",
            "badilisha jina",
        ]
    )


def create_docx_report(content: str, output_dir: Path, title: str = "Aggie Report") -> Path:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = output_dir / f"aggie-report-{timestamp}.docx"
    doc = Document()
    doc.add_heading(title, 0)
    for block in _split_blocks(content):
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif _looks_like_table(block):
            _add_docx_table(doc, block)
        else:
            doc.add_paragraph(block)
    doc.save(str(target))
    return target


def clean_docx_document(source: Path, output_dir: Path, instruction_text: str = "") -> Path:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = output_dir / f"{source.stem}-cleaned-{timestamp}.docx"
    doc = Document(str(source))
    title = _clean_heading(_extract_heading(instruction_text, source)).title()
    if doc.paragraphs:
        first = doc.paragraphs[0]
        if first.text.strip().lower() != title.lower():
            first.insert_paragraph_before(title, style="Title")
    else:
        doc.add_heading(title, 0)

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = None

    for table in doc.tables:
        table.style = "Table Grid"
        if table.rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.save(str(target))
    return target


def create_pdf_report(content: str, output_dir: Path, title: str = "Aggie Report") -> Path:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = output_dir / f"aggie-report-{timestamp}.pdf"
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for block in _split_blocks(content):
        if block.startswith("# "):
            story.append(Paragraph(block[2:].strip(), styles["Heading1"]))
        elif block.startswith("## "):
            story.append(Paragraph(block[3:].strip(), styles["Heading2"]))
        else:
            safe = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 8))
    SimpleDocTemplate(str(target), pagesize=A4).build(story)
    return target


def create_excel_from_text(content: str, output_dir: Path, title: str = "Organized Data") -> Path:
    from openpyxl import Workbook

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = output_dir / f"organized-data-{timestamp}.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Organized Data"
    rows = _text_to_rows(content)
    if not rows:
        rows = [["Content"], [content[:30000]]]
    for row in rows:
        sheet.append(row)
    _clean_sheet(sheet, title)
    _add_total_formulas(sheet)
    workbook.save(str(target))
    return target


def combined_attachment_text(attachments: list[dict[str, Any]]) -> str:
    chunks: list[str] = []
    for attachment in attachments:
        path = Path(attachment["path"])
        mime_type = attachment.get("mime_type")
        text = extract_document_text(path, mime_type).strip()
        if text:
            chunks.append(f"--- {attachment.get('filename') or path.name} ---\n{text}")
    return "\n\n".join(chunks)


def should_create_clean_excel(text: str, attachments: list[dict[str, Any]]) -> bool:
    lowered = text.lower()
    wants_edit = any(
        word in lowered
        for word in [
            "clean",
            "format",
            "edit",
            "arrange",
            "organize",
            "pang",
            "panga",
            "safisha",
            "tengeneza",
            "rekebisha",
            "weka sawa",
        ]
    )
    return wants_edit and has_excel_attachment(attachments)


def has_excel_attachment(attachments: list[dict[str, Any]]) -> bool:
    return any(_is_excel_attachment(item) for item in attachments)


def _is_excel_attachment(attachment: dict[str, Any]) -> bool:
    path = Path(attachment["path"])
    mime_type = attachment.get("mime_type") or ""
    filename = str(attachment.get("filename") or path.name).lower()
    return (
        path.suffix.lower() in {".xlsx", ".xlsm", ".xls"}
        or filename.endswith((".xlsx", ".xlsm", ".xls"))
        or mime_type
        in {
            XLSX_MIME_TYPE,
            XLS_MIME_TYPE,
            "application/vnd.ms-excel.sheet.macroenabled.12",
        }
    )


def _clean_sheet(sheet: Any, heading: str, light: bool = False) -> None:
    _delete_empty_rows(sheet)
    _delete_empty_columns(sheet)
    _add_heading(sheet, heading)
    sheet.freeze_panes = "A3"

    if not light:
        for row in sheet.iter_rows():
            for cell in row:
                cell.fill = PatternFill(fill_type=None)
                cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                cell.font = Font(
                    name=cell.font.name or "Calibri",
                    size=cell.font.sz or 11,
                    bold=cell.row in {1, 2},
                    italic=cell.font.italic,
                    color="000000",
                )
    else:
        for row in sheet.iter_rows(min_row=1, max_row=min(2, sheet.max_row)):
            for cell in row:
                if cell.row == 1:
                    continue
                cell.font = Font(
                    name=cell.font.name or "Calibri",
                    size=cell.font.sz or 11,
                    bold=True,
                    color="000000",
                )
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    _auto_size_columns(sheet)

    if sheet.max_row >= 2:
        if not light:
            for cell in sheet[2]:
                cell.font = copy(cell.font)
                cell.font = Font(name=cell.font.name or "Calibri", size=11, bold=True, color="000000")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.auto_filter.ref = f"A2:{get_column_letter(sheet.max_column)}{sheet.max_row}"
        _add_total_formulas(sheet)


def _extract_delete_column_requests(text: str) -> list[str]:
    if not text.strip():
        return []
    patterns = [
        r"(?:delete|remove|drop|futa|ondoa)\s+(?:the\s+)?(?:column|columns|col|safu)\s+(?:ya\s+|yenye\s+|called\s+|named\s+)?([^.;,\n]+)",
        r"(?:delete|remove|drop|futa|ondoa|toa)\s+([^.;,\n]+?)\s+(?:column|columns|col|safu)",
        r"(?:toa)\s+(?:column|columns|col|safu)\s+(?:ya\s+|yenye\s+)?([^.;,\n]+)",
    ]
    filler_words = {
        "the", "a", "an", "ya", "yenye", "called", "named", "from",
        "kwenye", "sheet", "file", "hii", "hiyo", "ile", "hizo",
        "zote", "all", "na", "and", "then", "halafu", "kisha",
    }
    values: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            raw = match.group(1)
            raw = re.split(
                r"\b(?:then|halafu|kisha|please|tafadhali|from|kwenye)\b",
                raw,
                flags=re.IGNORECASE,
            )[0]
            parts = re.split(r"\s*(?:,|/|\+|&|\band\b|\bna\b)\s*", raw, flags=re.IGNORECASE)
            for part in parts:
                cleaned = part.strip(" '\"`:-")
                if cleaned and cleaned.lower() not in filler_words and len(cleaned) > 1:
                    values.append(cleaned)
    column_mentions = re.findall(
        r"([a-zA-Z0-9 _-]{2,40}?)\s+(?:column|columns|col|safu)\b",
        text,
        flags=re.IGNORECASE,
    )
    for mention in column_mentions:
        cleaned = re.sub(
            r"^(?:hyo|hiyo|ile|the|ya|yenye|called|named)\s+",
            "",
            mention.strip(" '\"`:-"),
            flags=re.IGNORECASE,
        )
        cleaned = re.split(
            r"\b(?:delete|remove|drop|futa|ondoa|toa|and|na)\b",
            cleaned,
            flags=re.IGNORECASE,
        )[-1].strip(" '\"`:-")
        if cleaned and cleaned.lower() not in filler_words and len(cleaned) > 1:
            values.append(cleaned)
    return _unique_preserve_order(values)


def _find_header_row(sheet: Any) -> int:
    best_row = 1
    best_score = -1
    max_scan = min(sheet.max_row, 10)
    for row_idx in range(1, max_scan + 1):
        values = [
            sheet.cell(row_idx, col_idx).value
            for col_idx in range(1, sheet.max_column + 1)
        ]
        non_empty = [value for value in values if value not in (None, "")]
        text_count = sum(1 for value in non_empty if isinstance(value, str))
        score = len(non_empty) + text_count
        if score > best_score and len(non_empty) >= 2:
            best_row = row_idx
            best_score = score
    return best_row


def _normalized_headers(sheet: Any, header_row: int) -> list[str]:
    return [
        _normalize_header(str(sheet.cell(header_row, col_idx).value or ""))
        for col_idx in range(1, sheet.max_column + 1)
    ]


def _sheet_contains_value(sheet: Any, needle: str) -> bool:
    target = needle.strip().lower()
    for row in sheet.iter_rows():
        for cell in row:
            if str(cell.value or "").strip().lower() == target:
                return True
    return False


def _delete_requested_columns(sheet: Any, header_row: int, requested_columns: list[str]) -> list[str]:
    matches: list[tuple[int, str]] = []
    for request in requested_columns:
        col_idx = _match_column(sheet, header_row, request)
        if col_idx and all(existing_idx != col_idx for existing_idx, _ in matches):
            header = sheet.cell(header_row, col_idx).value
            label = str(header).strip() if header not in (None, "") else get_column_letter(col_idx)
            matches.append((col_idx, label))

    deleted_labels: list[str] = []
    for col_idx, label in sorted(matches, reverse=True):
        sheet.delete_cols(col_idx)
        deleted_labels.append(label)
    return list(reversed(deleted_labels))


def _keep_requested_columns(sheet: Any, header_row: int, requested_columns: list[str]) -> list[str]:
    keep_indexes: set[int] = set()
    kept_labels: list[str] = []
    for request in requested_columns:
        col_idx = _match_column(sheet, header_row, request)
        if col_idx:
            keep_indexes.add(col_idx)
            header = sheet.cell(header_row, col_idx).value
            kept_labels.append(str(header).strip() if header not in (None, "") else get_column_letter(col_idx))
    if not keep_indexes:
        return []
    for col_idx in range(sheet.max_column, 0, -1):
        if col_idx not in keep_indexes:
            sheet.delete_cols(col_idx)
    return _unique_preserve_order(kept_labels)


def _rename_requested_columns(sheet: Any, header_row: int, columns: Any) -> list[str]:
    if not isinstance(columns, dict):
        return []
    renamed: list[str] = []
    for old_name, new_name in columns.items():
        if not str(new_name).strip():
            continue
        col_idx = _match_column(sheet, header_row, str(old_name))
        if col_idx:
            old_value = sheet.cell(header_row, col_idx).value
            sheet.cell(header_row, col_idx).value = str(new_name).strip()
            renamed.append(f"{old_value or old_name} to {new_name}")
    return renamed


def _sort_sheet_by_column(sheet: Any, header_row: int, column: str, direction: str) -> str | None:
    col_idx = _match_column(sheet, header_row, column)
    if not col_idx:
        return None
    rows = list(sheet.iter_rows(min_row=header_row + 1, max_row=sheet.max_row, values_only=False))
    non_empty_rows = [
        [cell.value for cell in row]
        for row in rows
        if any(cell.value not in (None, "") for cell in row)
    ]
    if len(non_empty_rows) < 2:
        return None
    reverse = direction.lower().startswith("desc")
    value_index = col_idx - 1
    non_empty_rows.sort(key=lambda row: _sort_value(row[value_index] if value_index < len(row) else None), reverse=reverse)
    sheet.delete_rows(header_row + 1, sheet.max_row - header_row)
    for row in non_empty_rows:
        sheet.append(row)
    header = sheet.cell(header_row, col_idx).value
    return str(header or column)


def _add_product_summary(sheet: Any, header_row: int) -> bool:
    product_col = _find_product_column(sheet, header_row)
    if not product_col and header_row == 1 and sheet.max_row >= 2:
        header_row = 2
        product_col = _find_product_column(sheet, header_row)
    if not product_col:
        return False

    numeric_cols = _find_numeric_columns(sheet, header_row, exclude={product_col})
    count_only = not numeric_cols

    summary: dict[str, dict[int, float]] = {}
    counts: dict[str, int] = {}
    for row_idx in range(header_row + 1, sheet.max_row + 1):
        product = sheet.cell(row_idx, product_col).value
        if product in (None, ""):
            continue
        product_name = str(product).strip()
        if not product_name or product_name.upper() == "TOTAL":
            continue
        counts[product_name] = counts.get(product_name, 0) + 1
        if count_only:
            continue
        bucket = summary.setdefault(product_name, {col_idx: 0 for col_idx in numeric_cols})
        for col_idx in numeric_cols:
            number = _to_number(sheet.cell(row_idx, col_idx).value)
            if number is not None:
                bucket[col_idx] += number

    if not summary and not counts:
        return False

    start_row = sheet.max_row + 2
    end_col = 2 if count_only else 1 + len(numeric_cols)
    sheet.cell(start_row, 1).value = "PRODUCT SUMMARY"
    sheet.cell(start_row, 1).font = Font(name="Calibri", size=13, bold=True, color="000000")
    if end_col > 1:
        sheet.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=end_col)

    header_output_row = start_row + 1
    product_header = sheet.cell(header_row, product_col).value or "Product"
    sheet.cell(header_output_row, 1).value = product_header
    sheet.cell(header_output_row, 1).font = Font(name="Calibri", size=11, bold=True, color="000000")
    if count_only:
        sheet.cell(header_output_row, 2).value = "Count"
        sheet.cell(header_output_row, 2).font = Font(name="Calibri", size=11, bold=True, color="000000")
    else:
        for output_idx, col_idx in enumerate(numeric_cols, start=2):
            sheet.cell(header_output_row, output_idx).value = sheet.cell(header_row, col_idx).value
            sheet.cell(header_output_row, output_idx).font = Font(name="Calibri", size=11, bold=True, color="000000")

    row_idx = header_output_row + 1
    product_names = sorted(counts, key=str.lower)
    for product_name in product_names:
        sheet.cell(row_idx, 1).value = product_name
        if count_only:
            sheet.cell(row_idx, 2).value = counts[product_name]
        else:
            totals = summary.get(product_name, {})
            for output_idx, col_idx in enumerate(numeric_cols, start=2):
                sheet.cell(row_idx, output_idx).value = totals.get(col_idx, 0)
        row_idx += 1

    total_row = row_idx
    sheet.cell(total_row, 1).value = "TOTAL"
    sheet.cell(total_row, 1).font = Font(name="Calibri", size=11, bold=True, color="000000")
    for output_idx in range(2, end_col + 1):
        letter = get_column_letter(output_idx)
        cell = sheet.cell(total_row, output_idx)
        cell.value = f"=SUM({letter}{header_output_row + 1}:{letter}{total_row - 1})"
        cell.font = Font(name="Calibri", size=11, bold=True, color="000000")

    for row in sheet.iter_rows(min_row=start_row, max_row=total_row, min_col=1, max_col=end_col):
        for cell in row:
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    _auto_size_columns(sheet)
    return True


def _find_product_column(sheet: Any, header_row: int) -> int | None:
    preferred = [
        "product",
        "item",
        "items",
        "description",
        "sku",
        "name",
        "jina",
        "bidhaa",
        "material",
    ]
    for name in preferred:
        col_idx = _match_column(sheet, header_row, name)
        if col_idx:
            return col_idx
    semantic_col = _find_product_like_column(sheet, header_row)
    if semantic_col:
        return semantic_col
    return None


def _find_product_like_column(sheet: Any, header_row: int) -> int | None:
    best: tuple[int, int] | None = None
    for col_idx in range(1, sheet.max_column + 1):
        values: list[str] = []
        for row_idx in range(header_row, min(sheet.max_row, header_row + 120) + 1):
            value = sheet.cell(row_idx, col_idx).value
            if value in (None, ""):
                continue
            text = str(value).strip()
            if not text or _looks_like_phone_value(text) or _looks_like_date_value(text):
                continue
            if re.fullmatch(r"(?:DO|INV|PO)?\d{2,}[-/\w]*", text, flags=re.IGNORECASE):
                continue
            if _to_number(text) is not None:
                continue
            values.append(text.upper())
        if len(values) < 2:
            continue
        unique_count = len(set(values))
        repeat_count = len(values) - unique_count
        average_len = sum(len(value) for value in values) / len(values)
        score = repeat_count * 4 + min(unique_count, 20) + int(average_len >= 6)
        if best is None or score > best[0]:
            best = (score, col_idx)
    return best[1] if best and best[0] >= 8 else None


def _find_numeric_columns(sheet: Any, header_row: int, exclude: set[int]) -> list[int]:
    numeric_cols: list[int] = []
    for col_idx in range(1, sheet.max_column + 1):
        if col_idx in exclude:
            continue
        numeric_count = 0
        for row_idx in range(header_row + 1, min(sheet.max_row, header_row + 80) + 1):
            if _to_number(sheet.cell(row_idx, col_idx).value) is not None:
                numeric_count += 1
        if numeric_count >= 1:
            numeric_cols.append(col_idx)
    return numeric_cols[:8]


def _to_number(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        cleaned = value.replace(",", "").strip()
        if re.fullmatch(r"-?\d+(?:\.\d+)?", cleaned):
            return float(cleaned)
    return None


def _sort_value(value: Any) -> tuple[int, str]:
    if value in (None, ""):
        return (1, "")
    return (0, str(value).lower())


def _match_column(sheet: Any, header_row: int, request: str) -> int | None:
    cleaned_request = _normalize_header(request)
    if not cleaned_request:
        return None

    semantic_col = _semantic_column_match(sheet, header_row, cleaned_request)
    if semantic_col:
        return semantic_col

    if re.fullmatch(r"[a-zA-Z]{1,3}", request.strip()):
        col_idx = _column_letter_to_index(request.strip())
        if col_idx and col_idx <= sheet.max_column:
            return col_idx

    if request.strip().isdigit():
        col_idx = int(request.strip())
        if 1 <= col_idx <= sheet.max_column:
            return col_idx

    best: tuple[int, int] | None = None
    for col_idx in range(1, sheet.max_column + 1):
        value = sheet.cell(header_row, col_idx).value
        header = _normalize_header("" if value is None else str(value))
        if not header:
            continue
        score = 0
        if header == cleaned_request:
            score = 100
        elif cleaned_request in header or header in cleaned_request:
            score = 80
        elif _word_overlap(cleaned_request, header) >= 1:
            score = 60
        if score and (best is None or score > best[0]):
            best = (score, col_idx)
    return best[1] if best else None


def _semantic_column_match(sheet: Any, header_row: int, request: str) -> int | None:
    words = set(request.split())
    if words & {"simu", "phone", "mobile", "number", "namba", "contact", "contacts"}:
        return _find_phone_like_column(sheet, header_row)
    if words & {"quantity", "quantiti", "qty", "qnty", "idadi", "idad", "pcs", "piece", "pieces"}:
        return _find_quantity_like_column(sheet, header_row)
    return None


def _find_phone_like_column(sheet: Any, header_row: int) -> int | None:
    best: tuple[int, int] | None = None
    for col_idx in range(1, sheet.max_column + 1):
        hits = 0
        for row_idx in range(header_row + 1, min(sheet.max_row, header_row + 80) + 1):
            value = sheet.cell(row_idx, col_idx).value
            if _looks_like_phone_value(value):
                hits += 1
        if hits and (best is None or hits > best[0]):
            best = (hits, col_idx)
    return best[1] if best and best[0] >= 1 else None


def _find_quantity_like_column(sheet: Any, header_row: int) -> int | None:
    for name in ["qty", "quantity", "qnty", "pcs", "pieces", "idadi"]:
        for col_idx in range(1, sheet.max_column + 1):
            header = _normalize_header(str(sheet.cell(header_row, col_idx).value or ""))
            if name in header:
                return col_idx

    best: tuple[int, int] | None = None
    for col_idx in range(1, sheet.max_column + 1):
        numeric_count = 0
        integerish_count = 0
        for row_idx in range(header_row + 1, min(sheet.max_row, header_row + 80) + 1):
            number = _to_number(sheet.cell(row_idx, col_idx).value)
            if number is None:
                continue
            numeric_count += 1
            if float(number).is_integer() and 0 <= number <= 100000:
                integerish_count += 1
        score = integerish_count * 2 + numeric_count
        if score and (best is None or score > best[0]):
            best = (score, col_idx)
    return best[1] if best and best[0] >= 2 else None


def _looks_like_phone_value(value: Any) -> bool:
    if value in (None, ""):
        return False
    text = re.sub(r"\D", "", str(value))
    if len(text) < 7 or len(text) > 15:
        return False
    return text.startswith(("0", "255", "254", "256", "1", "7"))


def _looks_like_date_value(value: Any) -> bool:
    text = str(value).strip()
    return bool(
        re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}(?:\s+\d{1,2}:\d{2})?", text)
        or re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}(?:\s+\d{1,2}:\d{2}:\d{2})?", text)
    )


def _normalize_header(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _word_overlap(left: str, right: str) -> int:
    left_words = {word for word in left.split() if len(word) > 2}
    right_words = {word for word in right.split() if len(word) > 2}
    return len(left_words & right_words)


def _column_letter_to_index(value: str) -> int | None:
    result = 0
    for char in value.upper():
        if not ("A" <= char <= "Z"):
            return None
        result = result * 26 + (ord(char) - ord("A") + 1)
    return result or None


def _unique_preserve_order(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        key = value.lower()
        if key not in seen:
            seen.add(key)
            result.append(value)
    return result


def _normalize_excel_plan_actions(plan: dict[str, Any] | None) -> list[dict[str, Any]]:
    if not isinstance(plan, dict):
        return []
    actions = plan.get("actions")
    if not isinstance(actions, list):
        return []
    normalized: list[dict[str, Any]] = []
    allowed = {"delete_columns", "keep_columns", "rename_columns", "sort_by", "add_product_summary"}
    for action in actions:
        if isinstance(action, dict) and str(action.get("type", "")).lower().strip() in allowed:
            normalized.append(action)
    return normalized


def _string_list(value: Any) -> list[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return []


def _guess_header_index(rows: list[list[str]]) -> int | None:
    best_index: int | None = None
    best_score = -1
    for index, row in enumerate(rows[:5]):
        non_empty = [value for value in row if value.strip()]
        score = len(non_empty) + sum(1 for value in non_empty if not _looks_number_like(value))
        if len(non_empty) >= 2 and score > best_score:
            best_index = index
            best_score = score
    return best_index


def _looks_number_like(value: str) -> bool:
    return bool(re.fullmatch(r"[\d\s,.-]+", value.strip()))


def _extract_heading(instruction_text: str, source: Path) -> str:
    text = " ".join(instruction_text.split())
    patterns = [
        r"(?:heading|title|kichwa|header)\s*(?:ya|ni|is|:|-)?\s*['\"]?([^'\"]{3,80})",
        r"(?:weka|add|put)\s+(?:heading|title|kichwa)\s*(?:ya|as|:|-)?\s*['\"]?([^'\"]{3,80})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _clean_heading(match.group(1))
    return _clean_heading(source.stem.replace("_", " ").replace("-", " "))


def _clean_heading(value: str) -> str:
    heading = re.split(r"\b(?:na|and|then|kisha|halafu|please|tafadhali)\b", value, flags=re.IGNORECASE)[0]
    heading = re.sub(r"\s+", " ", heading).strip(" .:-_")
    return (heading or "Stock Report").upper()


def _add_heading(sheet: Any, heading: str) -> None:
    if sheet.max_row == 0 or sheet.max_column == 0:
        return
    first_row_values = [sheet.cell(1, col_idx).value for col_idx in range(1, sheet.max_column + 1)]
    already_has_heading = (
        len([value for value in first_row_values if value not in (None, "")]) == 1
        and sheet.max_column > 1
    )
    if not already_has_heading:
        sheet.insert_rows(1)

    end_column = max(sheet.max_column, 1)
    if end_column > 1:
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=end_column)
    cell = sheet.cell(1, 1)
    cell.value = heading
    cell.font = Font(name="Calibri", size=15, bold=True, color="000000")
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    sheet.row_dimensions[1].height = 24


def _add_heading_only(sheet: Any, heading: str) -> None:
    """Add a heading row and basic header formatting without resetting existing cell styles."""
    if sheet.max_row == 0 or sheet.max_column == 0:
        return
    first_row_values = [sheet.cell(1, col_idx).value for col_idx in range(1, sheet.max_column + 1)]
    already_has_heading = (
        len([value for value in first_row_values if value not in (None, "")]) == 1
        and sheet.max_column > 1
    )
    if not already_has_heading:
        sheet.insert_rows(1)

    end_column = max(sheet.max_column, 1)
    if end_column > 1:
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=end_column)
    cell = sheet.cell(1, 1)
    cell.value = heading
    cell.font = Font(name="Calibri", size=15, bold=True, color="000000")
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    sheet.row_dimensions[1].height = 24

    if sheet.max_row >= 2:
        for cell in sheet[2]:
            cell.font = Font(name="Calibri", size=11, bold=True, color="000000")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.auto_filter.ref = f"A2:{get_column_letter(sheet.max_column)}{sheet.max_row}"


def _delete_empty_rows(sheet: Any) -> None:
    for row_idx in range(sheet.max_row, 0, -1):
        if all(sheet.cell(row_idx, col_idx).value in (None, "") for col_idx in range(1, sheet.max_column + 1)):
            sheet.delete_rows(row_idx)


def _delete_empty_columns(sheet: Any) -> None:
    for col_idx in range(sheet.max_column, 0, -1):
        if all(sheet.cell(row_idx, col_idx).value in (None, "") for row_idx in range(1, sheet.max_row + 1)):
            sheet.delete_cols(col_idx)


def _auto_size_columns(sheet: Any) -> None:
    for col_idx in range(1, sheet.max_column + 1):
        letter = get_column_letter(col_idx)
        max_len = 10
        for cell in sheet[letter]:
            value = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, min(len(value), 45))
        sheet.column_dimensions[letter].width = max_len + 2


def _add_total_formulas(sheet: Any) -> None:
    if sheet.max_row < 3 or sheet.max_column < 2:
        return
    last_data_row = sheet.max_row
    if last_data_row > 3:
        first_cell = sheet.cell(last_data_row, 1).value
        if isinstance(first_cell, str) and first_cell.strip().upper() == "TOTAL":
            last_data_row -= 1
    if last_data_row < 3:
        return
    total_row = last_data_row + 1
    sheet.cell(total_row, 1).value = "TOTAL"
    sheet.cell(total_row, 1).font = Font(name="Calibri", size=11, bold=True, color="000000")
    for col_idx in range(2, sheet.max_column + 1):
        numeric_count = 0
        for row_idx in range(3, last_data_row + 1):
            value = sheet.cell(row_idx, col_idx).value
            if isinstance(value, (int, float)):
                numeric_count += 1
        if numeric_count >= 2:
            letter = get_column_letter(col_idx)
            cell = sheet.cell(total_row, col_idx)
            cell.value = f"=SUM({letter}3:{letter}{last_data_row})"
            cell.font = Font(name="Calibri", size=11, bold=True, color="000000")
            cell.alignment = Alignment(horizontal="right", vertical="center")


def _split_blocks(content: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", content.strip()) if block.strip()]


def _looks_like_table(block: str) -> bool:
    lines = [line for line in block.splitlines() if line.strip()]
    return len(lines) >= 2 and all("|" in line for line in lines[:2])


def _add_docx_table(doc: Document, block: str) -> None:
    rows = _text_to_rows(block)
    if not rows:
        doc.add_paragraph(block)
        return
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _text_to_rows(content: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in content.splitlines():
        cleaned = line.strip().strip("|")
        if not cleaned:
            continue
        if set(cleaned.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        if "|" in cleaned:
            row = [part.strip() for part in cleaned.split("|")]
        elif "," in cleaned:
            row = [part.strip() for part in next(csv.reader([cleaned]))]
        elif "\t" in cleaned:
            row = [part.strip() for part in cleaned.split("\t")]
        else:
            parts = re.split(r"\s{2,}", cleaned)
            row = parts if len(parts) > 1 else [cleaned]
        rows.append(row)
        if len(rows) >= 500:
            break
    return rows


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {index} ---\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_xlsx(path: Path) -> str:
    workbook = load_workbook(str(path), data_only=True, read_only=True)
    try:
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append(" | ".join(values))
                if len(rows) >= 200:
                    rows.append("[Sheet truncated after 200 non-empty rows]")
                    break
            if rows:
                chunks.append(f"--- Sheet: {sheet.title} ---\n" + "\n".join(rows))
        return "\n\n".join(chunks)
    finally:
        workbook.close()


def _extract_legacy_xls(path: Path) -> str:
    pd = _pandas()
    sheets = pd.read_excel(path, sheet_name=None, engine="xlrd", nrows=200)
    chunks: list[str] = []
    for sheet_name, frame in sheets.items():
        if frame.empty:
            continue
        rows = [" | ".join(map(str, frame.columns.tolist()))]
        rows.extend(" | ".join("" if value is None else str(value) for value in row) for row in frame.fillna("").values.tolist())
        chunks.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows[:201]))
    return "\n\n".join(chunks)


def _extract_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as file:
        reader = csv.reader(file)
        for row in reader:
            rows.append(" | ".join(row))
            if len(rows) >= 300:
                rows.append("[CSV truncated after 300 rows]")
                break
    return "\n".join(rows)


def _pandas() -> Any:
    import pandas as pd

    return pd
